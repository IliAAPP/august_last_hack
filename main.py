from typing import Dict, List
import uuid
import os
from docx import Document
from reportlab.pdfgen import canvas
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import nltk
from transformers import AutoTokenizer, AutoModelForSequenceClassification, pipeline
import whisper
import base64
from fastapi import File, UploadFile

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Или укажите конкретные домены
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

class DocumentCreate(BaseModel):
    type: str

class DocumentUpdate(BaseModel):
    fields: Dict[str, str]

class DocumentResponse(BaseModel):
    id: str
    type: str
    fields: Dict[str, str]

documents = {}

templates = {
    "contract": {
        "fields": {
            "client_name": "",
            "project_name": "",
            "start_date": "",
            "end_date": "",
            "contract_amount": "",
        },
        "title": "Договор на выполнение строительных работ",
        "formatting": {
            "client_name": "Клиент:",
            "project_name": "Проект:",
            "start_date": "Дата начала:",
            "end_date": "Дата окончания:",
            "contract_amount": "Сумма договора:"
        }
    },
    "estimate": {
        "fields": {
            "project_name": "",
            "materials": "",
            "labor_costs": "",
            "equipment_costs": "",
            "total_amount": "",
        },
        "title": "Смета на строительные работы",
        "formatting": {
            "project_name": "Проект:",
            "materials": "Материалы:",
            "labor_costs": "Трудозатраты:",
            "equipment_costs": "Оборудование:",
            "total_amount": "Итоговая сумма:"
        }
    },
    "work_order": {
        "fields": {
            "contractor_name": "",
            "work_description": "",
            "start_date": "",
            "end_date": "",
            "payment_terms": "",
        },
        "title": "Подрядный договор",
        "formatting": {
            "contractor_name": "Подрядчик:",
            "work_description": "Описание работы:",
            "start_date": "Дата начала:",
            "end_date": "Дата окончания:",
            "payment_terms": "Условия оплаты:"
        }
    },
}

@app.post("/documents", response_model=DocumentResponse)
async def create_document(document: DocumentCreate):
    doc_id = str(uuid.uuid4())
    if document.type not in templates:
        raise HTTPException(status_code=400, detail="Invalid document type")

    new_document = {
        "id": doc_id,
        "type": document.type,
        "fields": templates[document.type]["fields"].copy()
    }
    documents[doc_id] = new_document
    generate_initial_document(doc_id, new_document)
    return new_document

def generate_initial_document(doc_id: str, document: dict):
    doc_type = document["type"]
    if doc_type in ["contract", "work_order"]:
        generate_word_document(doc_id, document)
    elif doc_type == "estimate":
        generate_pdf_document(doc_id, document)

def generate_word_document(doc_id: str, document: dict):
    doc = Document()
    doc.add_heading(templates[document["type"]]["title"], 0)

    formatting = templates[document["type"]]["formatting"]
    for field, value in document["fields"].items():
        if value:  # Only add non-empty fields
            doc.add_paragraph(f"{formatting.get(field, field.replace('_', ' ').title())} {value}")

    filename = f"{doc_id}.docx"
    doc.save(os.path.join(UPLOAD_DIR, filename))

def generate_pdf_document(doc_id: str, document: dict):
    filename = f"{doc_id}.pdf"
    c = canvas.Canvas(os.path.join(UPLOAD_DIR, filename))
    c.setFont("Helvetica", 12)

    c.drawString(100, 750, templates[document["type"]]["title"])
    y = 730
    formatting = templates[document["type"]]["formatting"]
    for field, value in document["fields"].items():
        if value:  # Only add non-empty fields
            c.drawString(100, y, f"{formatting.get(field, field.replace('_', ' ').title())}: {value}")
            y -= 20

    c.save()

@app.get("/documents", response_model=List[DocumentResponse])
async def get_documents():
    return list(documents.values())

@app.get("/documents/{doc_id}", response_model=DocumentResponse)
async def get_document(doc_id: str):
    if doc_id not in documents:
        raise HTTPException(status_code=404, detail="Document not found")
    return documents[doc_id]

@app.put("/documents/{doc_id}", response_model=DocumentResponse)
async def update_document(doc_id: str, document: DocumentUpdate):
    if doc_id not in documents:
        raise HTTPException(status_code=404, detail="Document not found")

    documents[doc_id]["fields"].update(document.fields)

    # Generate and save the document
    doc_type = documents[doc_id]["type"]
    if doc_type in ["contract", "work_order"]:
        generate_word_document(doc_id, documents[doc_id])
    elif doc_type == "estimate":
        generate_pdf_document(doc_id, documents[doc_id])

    return documents[doc_id]

nltk.download('punkt')

model_name = "roberta-large-mnli"
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForSequenceClassification.from_pretrained(model_name)
classifier = pipeline("text-classification", model=model, tokenizer=tokenizer, return_all_scores=True)

class SentenceData(BaseModel):
    sentences: List[str]

@app.post("/check-contradictions")
async def check_contradictions(data: SentenceData):
    sentences = data.sentences
    if not sentences or len(sentences) < 2:
        raise HTTPException(status_code=400, detail="Должно быть как минимум два предложения для анализа")

    contradictions = 0
    total_pairs = 0

    for i in range(len(sentences)):
        for j in range(i + 1, len(sentences)):
            premise = sentences[i]
            hypothesis = sentences[j]
            # Прогоняем пару предложений через модель
            result = classifier(f"{premise} {hypothesis}", padding=True, truncation=True)

            # Ищем противоречие
            contradiction_score = None
            for item in result[0]:
                if item['label'] == 'CONTRADICTION':
                    contradiction_score = item['score']
                    break

            if contradiction_score and contradiction_score > 0.3:  # Порог для противоречия
                contradictions += 1
            total_pairs += 1

    percentage = (1 - contradictions / total_pairs) * 100 if total_pairs > 0 else 100

    return {
        "contradictions_found": contradictions,
        "total_pairs_checked": total_pairs,
        "validity_percentage": percentage
    }

whisper_model = whisper.load_model("base")

class AudioData(BaseModel):
    audio: str  # Base64 кодированный аудиофайл

@app.post("/transcribe-audio")
async def transcribe_audio(file: UploadFile = File(...)):
    try:
        # Save the uploaded file
        with open("temp_audio.wav", "wb") as f:
            content = await file.read()
            f.write(content)

        # Transcribe the audio file
        result = whisper_model.transcribe("temp_audio.wav")
        transcript = result["text"]

        # Check for specific keywords for home screen
        home_keywords = [
            "Главная", "Дом", "Home", "Календарь", "Календар", "Календарь!", "Календарь.", "Календарь.",
            "Home", "Calendar", "Calendars", "Calendars!", "Calendars.", "Calendars",
            "Kalender", "Kalender!", "Kalender.", "Kalender", "Kalenders", "Kalenders!", "Kalenders.",
            "Calendrier", "Calendrier!", "Calendrier.", "Calendriers", "Calendriers!", "Calendriers.",
            "Calendario", "Calendario!", "Calendario.", "Calendarios", "Calendarios!", "Calendarios.",
            "Calendario", "Calendario!", "Calendario.", "Calendari", "Calendari!", "Calendari.",
            "Calendário", "Calendário!", "Calendário.", "Calendários", "Calendários!", "Calendários.",
            "Календар", "Календарь", "Calendar", "Calend", "Calend", "Kalendario", "Kalender", "Calendario", "Kalimde.", "Kalimde!", "Kalimde", "home.", "home", "home!"
        ]

        if any(keyword in transcript for keyword in home_keywords):
            return {"redirect": "Home"}  # Redirect to 'Home' screen

        homescreen_words = [
            "Документ", "Шаблоны", "Документы", "Заполнить документы", "Document", "Документ!", "Документ.",  " Смята.", "Смята.", "Смята!"
        ]

        if any(keyword in transcript for keyword in homescreen_words):
            return {"redirect": "HomeScreen"}  # Redirect to 'Home' screen

        # Check for specific keywords for template selection screen
        template_keywords = ["Смета", "Договор", "Подряд", "Estimate", "Contract", "Work Order", "deal", "Deal", "deal.", "Deal.", "deal!", "Deal!"]
        if any(keyword in transcript for keyword in template_keywords):
            return {"redirect": "TemplateSelectionScreen"}  # Redirect to 'TemplateSelectionScreen'

        return {"transcript": transcript}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
